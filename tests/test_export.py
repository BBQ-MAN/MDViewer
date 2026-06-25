"""Word(.docx) 내보내기 코어 함수 단위 테스트 (Phase 10, 설계 10 §2/§10.1).

대상(PySide6 무의존 코어):
    - mdviewer.exporter.markdown_to_docx(markdown_text, out_path, base_dir, *, title=None) -> None
        · 변환은 절대 크래시하지 않는다(render 철학).
        · 최종 Document.save 의 I/O 실패(OSError)만 전파한다(write_markdown 철학).

경계면(설계 §9.2 A):
    - 패키지 루트 export 와 모듈 직접 import 가 동일 객체인지.
    - inspect.signature 가 계약 §2.1 과 정확히 일치하는지(keyword-only title 포함).

설계 문서: _workspace/10_export_feature_design.md
검증 포인트: §10.1 (기본 생성 / 매핑 대조 / 이미지 / 견고성 / I/O 전파 / title)

주의: 1x1 더미 PNG 는 python-docx 가 거부하므로, stdlib zlib+struct 로
완전한 IDAT 청크를 가진 유효 PNG(4x4 RGB)를 fixture 로 생성해 임베드를 검증한다.
"""

from __future__ import annotations

import inspect
import struct
import zlib
from pathlib import Path

import pytest

# 패키지 루트 export 와 모듈 직접 import 가 같은 객체인지 확인한다(경계면).
from mdviewer import markdown_to_docx as pkg_markdown_to_docx
from mdviewer.exporter import markdown_to_docx

# python-docx 는 테스트 전체의 전제. 없으면 전체 모듈 스킵(환경 결함을 숨기지 않음).
docx = pytest.importorskip("docx", reason="python-docx 미설치 — exporter 테스트 불가")
from docx import Document  # noqa: E402

# 대표 마크다운: 헤딩 h1~h3 / 단락 / 굵게·기울임·인라인코드 / 링크 / 중첩 목록 /
# 인용 / 표 / 코드펜스 / hr / task.
REPRESENTATIVE_MD = """# Title One

## Section Two

### Sub Three

Para with **bold** *em* `inline code` and [link](sub/page.md).

- a
- b
  - nested

1. one
2. two

> quote line

| H1 | H2 |
|----|----|
| a  | b  |

```python
def f():
    return 1
```

- [ ] task undone
- [x] task done

---
"""


# --------------------------------------------------------------------------- #
# 0. 경계면 — export 동일성 + 시그니처 shape 대조(설계 §10.3 / §9.2 A)
# --------------------------------------------------------------------------- #
def test_export_identity():
    """`from mdviewer import markdown_to_docx` 와 모듈 직접 import 가 동일 객체."""
    assert pkg_markdown_to_docx is markdown_to_docx


def test_signature_shape():
    """시그니처가 계약 §2.1 과 정확히 일치(keyword-only title, 반환 None)."""
    sig = inspect.signature(markdown_to_docx)
    params = list(sig.parameters.values())
    names = [p.name for p in params]
    assert names == ["markdown_text", "out_path", "base_dir", "title"]
    # 처음 3 개는 위치 가능, title 은 keyword-only.
    assert params[0].kind in (
        inspect.Parameter.POSITIONAL_OR_KEYWORD,
        inspect.Parameter.POSITIONAL_ONLY,
    )
    assert params[1].kind in (
        inspect.Parameter.POSITIONAL_OR_KEYWORD,
        inspect.Parameter.POSITIONAL_ONLY,
    )
    assert params[2].kind in (
        inspect.Parameter.POSITIONAL_OR_KEYWORD,
        inspect.Parameter.POSITIONAL_ONLY,
    )
    title_p = sig.parameters["title"]
    assert title_p.kind is inspect.Parameter.KEYWORD_ONLY
    assert title_p.default is None


# --------------------------------------------------------------------------- #
# 유효 PNG fixture (stdlib zlib + struct — 완전한 IDAT 청크)
# --------------------------------------------------------------------------- #
def _make_png(path: Path, w: int = 4, h: int = 4) -> None:
    """완전한 IHDR/IDAT/IEND 청크를 가진 유효 4x4 RGB PNG 를 생성한다."""

    def _chunk(typ: bytes, data: bytes) -> bytes:
        body = typ + data
        return (
            struct.pack(">I", len(data))
            + body
            + struct.pack(">I", zlib.crc32(body) & 0xFFFFFFFF)
        )

    sig = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0)  # 8-bit RGB
    raw = b""
    for y in range(h):
        raw += b"\x00"  # 필터 타입 None
        for x in range(w):
            v = 255 if (x + y) % 2 else 0
            raw += bytes([v, v, v])
    idat = zlib.compress(raw)
    path.write_bytes(
        sig + _chunk(b"IHDR", ihdr) + _chunk(b"IDAT", idat) + _chunk(b"IEND", b"")
    )


@pytest.fixture()
def valid_png(tmp_path: Path) -> Path:
    p = tmp_path / "pic.png"
    _make_png(p)
    assert p.stat().st_size > 0
    return p


# --------------------------------------------------------------------------- #
# 1. 기본 생성 — 파일 존재·크기·재오픈
# --------------------------------------------------------------------------- #
def test_basic_creation(tmp_path: Path):
    out = tmp_path / "out.docx"
    markdown_to_docx(REPRESENTATIVE_MD, out, tmp_path)
    assert out.exists()
    assert out.stat().st_size > 0

    doc = Document(str(out))
    assert len(doc.paragraphs) > 0
    assert len(doc.tables) == 1

    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Title One" in all_text
    assert "Section Two" in all_text
    assert "Sub Three" in all_text
    assert "quote line" in all_text


def test_returns_none(tmp_path: Path):
    """반환 shape: 항상 None."""
    out = tmp_path / "n.docx"
    rv = markdown_to_docx("# hi", out, tmp_path)
    assert rv is None


# --------------------------------------------------------------------------- #
# 2. 매핑 대조 (설계 §10.1)
# --------------------------------------------------------------------------- #
def test_heading_style(tmp_path: Path):
    """h1 → Heading 1 스타일."""
    out = tmp_path / "h.docx"
    markdown_to_docx("# Title One\n\n## Two\n", out, tmp_path)
    doc = Document(str(out))
    first = doc.paragraphs[0]
    assert first.text == "Title One"
    assert first.style.name == "Heading 1"
    # h2 도 어딘가에 Heading 2 로.
    assert any(
        p.text == "Two" and p.style.name == "Heading 2" for p in doc.paragraphs
    )


def test_bold_run(tmp_path: Path):
    """굵게 → run.bold is True."""
    out = tmp_path / "b.docx"
    markdown_to_docx("Para with **bold** here.", out, tmp_path)
    doc = Document(str(out))
    bold_runs = [r for p in doc.paragraphs for r in p.runs if r.bold]
    assert bold_runs, "굵게 run 이 없음"
    assert any("bold" in r.text for r in bold_runs)
    assert bold_runs[0].bold is True


def test_italic_run(tmp_path: Path):
    out = tmp_path / "i.docx"
    markdown_to_docx("Para with *em* here.", out, tmp_path)
    doc = Document(str(out))
    italic_runs = [r for p in doc.paragraphs for r in p.runs if r.italic]
    assert italic_runs, "기울임 run 이 없음"
    assert any("em" in r.text for r in italic_runs)


def test_table_mapping(tmp_path: Path):
    """표: len(doc.tables)==1, 셀 텍스트 일치."""
    md = "| H1 | H2 |\n|----|----|\n| a  | b  |\n"
    out = tmp_path / "t.docx"
    markdown_to_docx(md, out, tmp_path)
    doc = Document(str(out))
    assert len(doc.tables) == 1
    table = doc.tables[0]
    # 헤더 + 1 데이터 행, 2 열.
    assert len(table.rows) == 2
    assert len(table.columns) == 2
    assert table.rows[0].cells[0].text == "H1"
    assert table.rows[0].cells[1].text == "H2"
    assert table.rows[1].cells[0].text == "a"
    assert table.rows[1].cells[1].text == "b"


def test_code_block_preserved(tmp_path: Path):
    """코드블록 텍스트(개행 포함) 보존."""
    md = "```python\ndef f():\n    return 1\n```\n"
    out = tmp_path / "c.docx"
    markdown_to_docx(md, out, tmp_path)
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    # 각 라인이 별도 monospace 단락으로 보존되어야 한다.
    assert any(t == "def f():" for t in texts), texts
    assert any(t == "    return 1" for t in texts), texts
    # monospace(Consolas) 적용 확인.
    mono = [
        r
        for p in doc.paragraphs
        for r in p.runs
        if r.font.name == "Consolas" and "def f()" in r.text
    ]
    assert mono, "코드블록 run 이 Consolas 가 아님"


def test_inline_code(tmp_path: Path):
    """인라인 코드 → monospace run."""
    out = tmp_path / "ic.docx"
    markdown_to_docx("Use `inline code` now.", out, tmp_path)
    doc = Document(str(out))
    mono = [
        r
        for p in doc.paragraphs
        for r in p.runs
        if r.font.name == "Consolas" and "inline code" in r.text
    ]
    assert mono, "인라인 코드가 monospace 가 아님"


def test_task_list_prefix(tmp_path: Path):
    """task 항목 → ☑/☐ 접두."""
    md = "- [ ] task undone\n- [x] task done\n"
    out = tmp_path / "task.docx"
    markdown_to_docx(md, out, tmp_path)
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert any(t == "☐ task undone" for t in texts), texts  # ☐
    assert any(t == "☑ task done" for t in texts), texts    # ☑


def test_nested_list_no_loss(tmp_path: Path):
    """중첩 목록 → 크래시 없음 + 모든 항목 텍스트 보존 + List 스타일 사용.

    구조(중첩 항목이 별도 단락으로 분리)는 BUG-01(아래 xfail)로 추적한다.
    여기서는 변환이 크래시하지 않고 텍스트가 유실되지 않음만 보증한다.
    """
    md = "- a\n- b\n  - nested\n"
    out = tmp_path / "nl.docx"
    markdown_to_docx(md, out, tmp_path)
    doc = Document(str(out))
    full = "\n".join(p.text for p in doc.paragraphs)
    # 텍스트 유실 없음.
    assert "a" in full
    assert "b" in full
    assert "nested" in full
    styles = [p.style.name for p in doc.paragraphs if p.text.strip()]
    assert any("List" in s for s in styles), "List 스타일 단락이 없음"


def test_nested_list_structure(tmp_path: Path):
    """중첩 목록: 부모 항목과 자식 항목이 각각 별도 단락이어야 한다(정상 구조).

    BUG-01 수정 후 회귀 가드. exporter 는 중첩 <ul>/<ol> 진입 시 부모 <li> 인라인을
    먼저 flush 하므로, 부모 'b' 와 자식 'nested' 가 **각각 독립 단락**으로 분리되고
    잉여 개행이 끼지 않는다(자식은 더 깊은 List 스타일).
    """
    md = "- a\n- b\n  - nested\n"
    out = tmp_path / "nl.docx"
    markdown_to_docx(md, out, tmp_path)
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    # 부모 항목 'b' 가 'nested' 와 분리된 자체 단락이어야 한다(병합·잉여개행 없음).
    assert "b" in texts, texts
    assert "nested" in texts, texts
    # 'b' 와 'nested' 가 한 단락으로 병합되지 않았는지(BUG-01 의 핵심 증상).
    assert not any("nested" in t and "b" in t.split() for t in texts), texts
    # 자식 항목은 부모보다 깊은 들여쓰기(List ... 2) 스타일이어야 한다.
    styles = {p.text: (p.style.name if p.style else "") for p in doc.paragraphs if p.text.strip()}
    assert "2" in styles.get("nested", ""), styles


# --------------------------------------------------------------------------- #
# 3. 이미지 (설계 §10.1)
# --------------------------------------------------------------------------- #
def test_image_embed_valid_png(valid_png: Path, tmp_path: Path):
    """유효 로컬 PNG → 임베드(inline_shapes >= 1)."""
    # base_dir = valid_png.parent 이므로 상대경로로 참조.
    md = "![alt text](pic.png)\n"
    out = tmp_path / "img.docx"
    markdown_to_docx(md, out, valid_png.parent)
    doc = Document(str(out))
    assert len(doc.inline_shapes) >= 1


def test_image_missing_file_alt_fallback(tmp_path: Path):
    """존재하지 않는 file:// → 예외 없이 alt 폴백, 임베드 없음."""
    md = "![missing alt](nope.png)\n"
    out = tmp_path / "miss.docx"
    markdown_to_docx(md, out, tmp_path)  # 예외 없어야 함
    doc = Document(str(out))
    assert len(doc.inline_shapes) == 0
    assert any("missing alt" in p.text for p in doc.paragraphs)


def test_image_remote_alt_fallback(tmp_path: Path):
    """원격 http:// → 다운로드 없이 alt 폴백, 임베드 없음."""
    md = "![remote alt](http://example.com/x.png)\n"
    out = tmp_path / "rem.docx"
    markdown_to_docx(md, out, tmp_path)  # 다운로드/예외 없어야 함
    doc = Document(str(out))
    assert len(doc.inline_shapes) == 0
    assert any("remote alt" in p.text for p in doc.paragraphs)


# --------------------------------------------------------------------------- #
# 4. 견고성 (비전파 — 설계 §10.1)
# --------------------------------------------------------------------------- #
@pytest.mark.parametrize(
    "value",
    ["", "   ", "\n\n", None, "<b>unclosed <i>nested </b> text", "```\nunclosed code"],
    ids=["empty", "spaces", "newlines", "none", "broken-html", "unclosed-fence"],
)
def test_robustness_no_exception(value, tmp_path: Path):
    """빈/None/깨진 입력 → 예외 없이 .docx 저장."""
    out = tmp_path / "robust.docx"
    markdown_to_docx(value, out, tmp_path)  # 절대 예외 없어야 함
    assert out.exists()
    assert out.stat().st_size > 0
    # 재오픈 가능(유효 docx) 확인.
    Document(str(out))


def test_empty_creates_parent_dir(tmp_path: Path):
    """부모 디렉터리 없으면 생성(write_markdown 대칭)."""
    out = tmp_path / "sub" / "deep" / "out.docx"
    assert not out.parent.exists()
    markdown_to_docx("# hi", out, tmp_path)
    assert out.exists()


# --------------------------------------------------------------------------- #
# 5. I/O 전파 (OSError 전파 — 설계 §10.1 / §2.2)
# --------------------------------------------------------------------------- #
def test_io_error_propagates_directory_out_path(tmp_path: Path):
    """out_path 가 존재하는 디렉터리면 저장 실패 → OSError 전파."""
    dir_as_out = tmp_path / "adir"
    dir_as_out.mkdir()
    with pytest.raises(OSError):
        markdown_to_docx("# hi", dir_as_out, tmp_path)


# --------------------------------------------------------------------------- #
# 6. title (core_properties — 설계 §10.1)
# --------------------------------------------------------------------------- #
def test_title_core_property(tmp_path: Path):
    """title='X' → Document(out).core_properties.title == 'X'."""
    out = tmp_path / "title.docx"
    markdown_to_docx("# Body Heading\n", out, tmp_path, title="My Document Title")
    doc = Document(str(out))
    assert doc.core_properties.title == "My Document Title"


def test_title_omitted_no_crash(tmp_path: Path):
    """title 미지정 → 정상 저장(코어 속성 title 은 기본값)."""
    out = tmp_path / "notitle.docx"
    markdown_to_docx("# Body\n", out, tmp_path)
    doc = Document(str(out))
    # 본문 h1 은 별도 제목 단락이 아니라 본문에 들어가야 한다(설계 §2.1).
    assert doc.paragraphs[0].text == "Body"
