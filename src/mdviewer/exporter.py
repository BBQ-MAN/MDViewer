"""마크다운 → Word(.docx) 내보내기 (비-GUI 코어).

이 모듈은 **PySide6 에 의존하지 않는다.** 클립보드·파일 대화상자·페이지 인쇄는
전적으로 UI 계층의 책임이며, core 는 "이미 추출된 마크다운 문자열 + 출력 경로 +
base_dir" 만 받아 결정적인 .docx 산출물을 만든다. 이 경계가 GUI 없는 단위 테스트를
가능하게 한다(설계 §2 / 10_export_feature_design.md).

변환 흐름:
    markdown_text
      └─ renderer.render(markdown_text, base_dir).html   (본문 HTML)
           └─ stdlib html.parser 로 walk            (bs4/lxml 직접 파싱 금지)
                └─ python-docx 요소로 매핑 → Document.save(out_path)

예외 정책(설계 §2.1 / §2.2 / §2.7):
    - **변환 자체는 절대 크래시하지 않는다**(renderer.render 의 "어떤 입력에도 안전"
      철학). 깨진/빈/비정상 HTML, 알 수 없는 태그, 깨진 이미지에도 예외를 던지지
      않고 최선의 docx 를 만들어 저장한다.
    - **단, 최종 ``Document.save(out_path)`` 의 I/O 실패(OSError)만 전파**한다
      (write_markdown 과 대칭 — "저장 실패"는 사용자가 반드시 알아야 하는 사건).

python-docx(및 그 트랜지티브 의존 lxml)는 **함수 내부에서 지연 import** 한다.
이렇게 하면 내보내기를 쓰지 않는 한 lxml C 확장이 로드되지 않아 렌더 핫패스에
영향을 주지 않는다(설계 §6.1).
"""

from __future__ import annotations

import urllib.parse
import urllib.request
from html.parser import HTMLParser
from pathlib import Path
from typing import TYPE_CHECKING

from .renderer import render

if TYPE_CHECKING:  # 타입 힌트 전용(런타임 import 회피 — lxml 지연 로딩 유지).
    from docx.document import Document as _DocxDocument
    from docx.text.paragraph import Paragraph as _DocxParagraph
    from docx.text.run import Run as _DocxRun

__all__ = ["markdown_to_docx"]

# 본문 폭(A4 - 좌우 여백 ≈ 6.3in). 이미지 폭 클램프 기준(설계 §2.5).
_BODY_WIDTH_INCHES = 6.3

# 코드/인라인코드 monospace 글꼴·크기(설계 §2.4).
_MONO_FONT = "Consolas"
_CODE_FONT_SIZE_PT = 10

# 블록으로 취급하는 컨테이너(이 태그를 만나면 인라인 누적을 flush 한다).
_BLOCK_TAGS = frozenset(
    {
        "h1", "h2", "h3", "h4", "h5", "h6",
        "p", "ul", "ol", "li", "blockquote",
        "pre", "table", "thead", "tbody", "tr", "th", "td",
        "hr", "div", "img",
    }
)


# ---------------------------------------------------------------------------
# OXML 헬퍼 (python-docx 직접 API 가 없는 부분)
# ---------------------------------------------------------------------------
def _set_run_shading(run: "_DocxRun", fill: str) -> None:
    """run 에 옅은 배경 음영(``w:shd``)을 준다(인라인 코드/코드블록용, 선택적).

    python-docx 에 run 음영 고수준 API 가 없으므로 OXML 로 직접 주입한다.
    실패해도 조용히 무시한다(음영은 어디까지나 장식 — monospace 만으로 충분).
    """
    try:
        from docx.oxml.ns import qn

        rpr = run._element.get_or_add_rPr()
        shd = rpr.find(qn("w:shd"))
        if shd is None:
            shd = rpr.makeelement(qn("w:shd"), {})
            rpr.append(shd)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
    except Exception:
        pass


def _add_hyperlink(paragraph: "_DocxParagraph", url: str, text: str) -> bool:
    """단락에 클릭 가능한 ``w:hyperlink`` 를 추가한다(정식 하이퍼링크).

    python-docx 에 고수준 API 가 없어 관계(relationship) + OXML 로 직접 구성한다.
    파란 밑줄 스타일의 run 으로 표시한다. 실패하면 ``False`` 를 돌려 호출 측이
    일반 텍스트 run 으로 폴백하게 한다(설계 §2.6 — 변환은 절대 크래시 금지).

    Returns:
        bool: 하이퍼링크 추가 성공 여부.
    """
    try:
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.oxml.ns import qn
        from docx.oxml.shared import OxmlElement

        part = paragraph.part
        r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")

        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")
        r_pr.append(color)

        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        r_pr.append(underline)

        new_run.append(r_pr)

        text_el = OxmlElement("w:t")
        text_el.text = text or url
        text_el.set(qn("xml:space"), "preserve")
        new_run.append(text_el)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return True
    except Exception:
        return False


# ---------------------------------------------------------------------------
# 인라인 서식 스택 프레임
# ---------------------------------------------------------------------------
class _Inline:
    """walk 중 만난 인라인 텍스트 조각 + 활성 서식 스냅샷.

    text 노드를 만나면 현재 서식 집합(bold/italic/strike/code)과 활성 링크 href 를
    함께 보관해 두었다가, 블록 종료 시점에 한꺼번에 run 으로 flush 한다.
    """

    __slots__ = ("text", "bold", "italic", "strike", "code", "href", "is_break")

    def __init__(
        self,
        text: str = "",
        *,
        bold: bool = False,
        italic: bool = False,
        strike: bool = False,
        code: bool = False,
        href: str | None = None,
        is_break: bool = False,
    ) -> None:
        self.text = text
        self.bold = bold
        self.italic = italic
        self.strike = strike
        self.code = code
        self.href = href
        self.is_break = is_break


# ---------------------------------------------------------------------------
# HTML → docx 빌더 (stdlib html.parser 기반 walk)
# ---------------------------------------------------------------------------
class _DocxBuilder(HTMLParser):
    """renderer 본문 HTML 을 walk 하며 python-docx 요소를 조립한다.

    설계 §2.3 매핑표를 따른다. 어떤 단계의 내부 예외도 삼켜서 변환을 계속하며
    (부분 결과라도 보존), 텍스트 유실을 최소화한다.
    """

    def __init__(self, doc: "_DocxDocument", base_dir: Path) -> None:
        super().__init__(convert_charrefs=True)
        self._doc = doc
        self._base_dir = base_dir

        # 인라인 서식 활성 카운트(중첩 대응: <strong><strong>x</strong></strong>).
        self._bold = 0
        self._italic = 0
        self._strike = 0
        self._code = 0
        self._href: str | None = None

        # 현재 누적 중인 인라인 조각들(블록 종료 시 flush).
        self._inlines: list[_Inline] = []

        # 리스트 중첩 스택: ("ul"|"ol", depth). depth 1~3.
        self._list_stack: list[str] = []

        # 블록 컨텍스트 스택(p/li/blockquote/heading/th/td 등).
        # 어떤 단락에 인라인을 모으는지 판단하는 데 사용.
        self._ctx: list[str] = []

        # 헤딩 레벨(heading 컨텍스트 진입 시 설정).
        self._heading_level = 0

        # blockquote 중첩 깊이(내부 p 를 Quote 스타일로).
        self._in_blockquote = 0

        # 코드블록(<pre>) 처리: 원시 텍스트를 그대로 누적.
        self._in_pre = 0
        self._pre_text: list[str] = []
        self._pre_lang: str | None = None

        # 테이블 누적 상태.
        self._table_stack: list[dict] = []
        # 현재 셀 텍스트 누적(테이블 셀 안의 인라인은 단순 텍스트로 평탄화).
        self._cell_inlines: list[_Inline] | None = None

        # task-list 항목 여부 + 체크 상태(li 시작 시 판정).
        self._li_task: list[bool | None] = []  # None=일반 li, True/False=task

    # -- 텍스트 누적 -------------------------------------------------------
    def _push_text(self, text: str) -> None:
        """현재 서식으로 text 조각을 인라인 버퍼(또는 셀 버퍼)에 추가."""
        if not text:
            return
        frag = _Inline(
            text,
            bold=self._bold > 0,
            italic=self._italic > 0,
            strike=self._strike > 0,
            code=self._code > 0,
            href=self._href,
        )
        if self._cell_inlines is not None:
            self._cell_inlines.append(frag)
        else:
            self._inlines.append(frag)

    # -- run 생성 ----------------------------------------------------------
    def _emit_runs(self, paragraph: "_DocxParagraph", inlines: list[_Inline]) -> None:
        """인라인 조각 리스트를 paragraph 의 run/하이퍼링크로 실체화한다."""
        for frag in inlines:
            try:
                if frag.is_break:
                    if paragraph.runs:
                        paragraph.runs[-1].add_break()
                    else:
                        paragraph.add_run().add_break()
                    continue
                if not frag.text:
                    continue
                # 하이퍼링크(정식 시도 → 실패 시 텍스트 run 폴백).
                if frag.href:
                    if _add_hyperlink(paragraph, frag.href, frag.text):
                        continue
                    # 폴백: 텍스트 run + 괄호 URL.
                    run = paragraph.add_run(frag.text)
                    self._apply_run_format(run, frag)
                    tail = paragraph.add_run(f" ({frag.href})")
                    tail.italic = True
                    continue
                run = paragraph.add_run(frag.text)
                self._apply_run_format(run, frag)
            except Exception:
                # 한 조각 실패가 전체를 깨지 않게.
                try:
                    paragraph.add_run(frag.text or "")
                except Exception:
                    pass

    def _apply_run_format(self, run: "_DocxRun", frag: _Inline) -> None:
        """run 에 서식(bold/italic/strike/monospace 코드)을 적용."""
        try:
            if frag.bold:
                run.bold = True
            if frag.italic:
                run.italic = True
            if frag.strike:
                run.font.strike = True
            if frag.code:
                run.font.name = _MONO_FONT
                run.font.size = _pt(_CODE_FONT_SIZE_PT)
                _set_run_shading(run, "F2F2F2")
        except Exception:
            pass

    # -- 인라인 버퍼 flush --------------------------------------------------
    def _flush_block_paragraph(self, style: str | None = None) -> None:
        """누적 인라인을 새 단락으로 flush 한다(p/heading 외 일반 블록용)."""
        inlines = self._inlines
        self._inlines = []
        # 공백만이면 빈 단락을 만들지 않는다(과도한 빈 줄 방지).
        if not _has_visible(inlines):
            return
        try:
            para = self._doc.add_paragraph(style=style) if style else self._doc.add_paragraph()
        except Exception:
            para = self._doc.add_paragraph()
        self._emit_runs(para, inlines)

    # -- HTMLParser 콜백 ---------------------------------------------------
    def handle_starttag(self, tag: str, attrs_list) -> None:
        try:
            self._start(tag, dict(attrs_list))
        except Exception:
            pass

    def handle_startendtag(self, tag: str, attrs_list) -> None:
        # 자기완결 태그(<br/>, <hr/>, <img/>, <input/>).
        try:
            attrs = dict(attrs_list)
            if tag == "br":
                self._handle_br()
            elif tag == "hr":
                self._handle_hr()
            elif tag == "img":
                self._handle_img(attrs)
            elif tag == "input":
                self._handle_input(attrs)
            else:
                self._start(tag, attrs)
                self._end(tag)
        except Exception:
            pass

    def handle_endtag(self, tag: str) -> None:
        try:
            self._end(tag)
        except Exception:
            pass

    def handle_data(self, data: str) -> None:
        try:
            if self._in_pre:
                self._pre_text.append(data)
                return
            # 블록 사이의 순수 개행/들여쓰기 공백은 버린다(인라인 컨텍스트 밖).
            if not self._ctx and not data.strip():
                return
            self._push_text(data)
        except Exception:
            pass

    # -- 시작 태그 처리 -----------------------------------------------------
    def _start(self, tag: str, attrs: dict) -> None:
        if tag in ("strong", "b"):
            self._bold += 1
        elif tag in ("em", "i"):
            self._italic += 1
        elif tag in ("s", "del", "strike"):
            self._strike += 1
        elif tag == "code":
            # <pre><code> 의 code 는 코드블록(아래 pre 가 처리) — 인라인만 카운트.
            if not self._in_pre:
                self._code += 1
        elif tag == "a":
            href = attrs.get("href")
            self._href = href if isinstance(href, str) else None
        elif tag == "br":
            self._handle_br()
        elif tag == "img":
            self._handle_img(attrs)
        elif tag == "input":
            self._handle_input(attrs)
        elif tag == "hr":
            self._handle_hr()
        elif tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self._ctx.append(tag)
            self._heading_level = int(tag[1])
        elif tag == "p":
            self._ctx.append("p")
        elif tag == "blockquote":
            self._in_blockquote += 1
            self._ctx.append("blockquote")
        elif tag in ("ul", "ol"):
            # 부모 <li> 가 아직 열린 채(자식 목록 직전)면, 부모 항목의 누적 인라인을
            # 먼저 flush 한다. 안 그러면 부모 텍스트가 첫 자식 항목 단락에 병합되고
            # 자식 깊이 스타일(List Bullet 2 등)로 잘못 렌더된다(BUG-01).
            # 이 시점의 _list_stack/_li_task[-1] 은 아직 '부모' 레벨을 가리키므로,
            # _flush_list_item 이 부모 깊이/종류/태스크 상태로 올바르게 emit 한다.
            # flush 후 _inlines 가 비므로 이후 부모 </li> 의 재flush 는 no-op.
            if self._ctx and self._ctx[-1] == "li":
                self._flush_list_item()
            self._list_stack.append(tag)
        elif tag == "li":
            self._ctx.append("li")
            # task-list 항목인지 판정(class 에 task-list-item).
            cls = attrs.get("class") or ""
            self._li_task.append(False if "task-list-item" in cls else None)
        elif tag == "pre":
            self._in_pre += 1
            self._pre_text = []
            self._pre_lang = None
        elif tag == "table":
            self._table_stack.append({"rows": [], "cols": 0})
        elif tag == "tr":
            if self._table_stack:
                self._table_stack[-1].setdefault("_cur", [])
                self._table_stack[-1]["_cur"] = []
        elif tag in ("th", "td"):
            # 셀 텍스트 누적 시작.
            self._cell_inlines = []

    # -- 종료 태그 처리 -----------------------------------------------------
    def _end(self, tag: str) -> None:
        if tag in ("strong", "b"):
            self._bold = max(0, self._bold - 1)
        elif tag in ("em", "i"):
            self._italic = max(0, self._italic - 1)
        elif tag in ("s", "del", "strike"):
            self._strike = max(0, self._strike - 1)
        elif tag == "code":
            if not self._in_pre:
                self._code = max(0, self._code - 1)
        elif tag == "a":
            self._href = None
        elif tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self._flush_heading()
            if self._ctx and self._ctx[-1] == tag:
                self._ctx.pop()
            self._heading_level = 0
        elif tag == "p":
            self._flush_paragraph()
            if self._ctx and self._ctx[-1] == "p":
                self._ctx.pop()
        elif tag == "blockquote":
            self._in_blockquote = max(0, self._in_blockquote - 1)
            if self._ctx and self._ctx[-1] == "blockquote":
                self._ctx.pop()
        elif tag in ("ul", "ol"):
            if self._list_stack:
                self._list_stack.pop()
        elif tag == "li":
            self._flush_list_item()
            if self._ctx and self._ctx[-1] == "li":
                self._ctx.pop()
            if self._li_task:
                self._li_task.pop()
        elif tag == "pre":
            self._flush_code_block()
            self._in_pre = max(0, self._in_pre - 1)
        elif tag == "table":
            self._flush_table()
        elif tag == "tr":
            if self._table_stack and "_cur" in self._table_stack[-1]:
                row = self._table_stack[-1]["_cur"]
                self._table_stack[-1]["rows"].append(row)
                self._table_stack[-1]["cols"] = max(
                    self._table_stack[-1]["cols"], len(row)
                )
                self._table_stack[-1]["_cur"] = []
        elif tag in ("th", "td"):
            # 셀 텍스트를 평탄화해 현재 행에 추가.
            text = _inline_plain(self._cell_inlines or [])
            self._cell_inlines = None
            if self._table_stack:
                self._table_stack[-1].setdefault("_cur", [])
                self._table_stack[-1]["_cur"].append(text)

    # -- 블록 flush 구현 ---------------------------------------------------
    def _flush_heading(self) -> None:
        inlines = self._inlines
        self._inlines = []
        level = self._heading_level or 1
        text = _inline_plain(inlines)
        try:
            # add_heading 은 단일 텍스트만 받으므로 평문으로(헤딩 내 서식은 드묾).
            self._doc.add_heading(text, level=min(max(level, 0), 9))
        except Exception:
            try:
                self._doc.add_paragraph(text)
            except Exception:
                pass

    def _flush_paragraph(self) -> None:
        # blockquote 내부 p 는 Quote 스타일로.
        if self._in_blockquote > 0:
            self._flush_block_paragraph(style=_safe_style(self._doc, "Intense Quote", "Quote"))
            return
        # 리스트 항목 내부의 p 가 아니면 일반 단락.
        self._flush_block_paragraph()

    def _flush_list_item(self) -> None:
        inlines = self._inlines
        self._inlines = []
        # HTML 직렬화가 항목/자식목록 경계에 끼워 넣는 구조적 공백(개행/들여쓰기)을
        # 항목 경계에서 다듬는다(BUG-01 부수: 'Item B\n' / '\nNested' 잔여 개행 제거).
        # 첫/마지막 텍스트 조각의 바깥쪽 공백만 트림 — 내부 간격·인라인 서식은 보존.
        _trim_boundary_ws(inlines)
        if not _has_visible(inlines) and not (self._li_task and self._li_task[-1] is not None):
            return
        depth = min(len(self._list_stack), 3) or 1
        kind = self._list_stack[-1] if self._list_stack else "ul"
        task_state = self._li_task[-1] if self._li_task else None

        if task_state is not None:
            # task-list 항목: ☑/☐ 접두 + List Bullet.
            # renderer 는 <input> 뒤 텍스트에 선행 공백을 남기므로(예: " task"),
            # 첫 텍스트 조각의 선행 공백 1개를 다듬어 "☑ task" 로 깔끔히 만든다.
            for frag in inlines:
                if frag.is_break:
                    break
                if frag.text:
                    frag.text = frag.text.lstrip(" ")
                    break
            prefix = "☑ " if task_state else "☐ "  # ☑ / ☐
            style = _safe_style(self._doc, "List Bullet", None)
            try:
                para = self._doc.add_paragraph(style=style) if style else self._doc.add_paragraph()
            except Exception:
                para = self._doc.add_paragraph()
            para.add_run(prefix)
            self._emit_runs(para, inlines)
            return

        if kind == "ol":
            base = "List Number"
        else:
            base = "List Bullet"
        style_name = base if depth <= 1 else f"{base} {depth}"
        style = _safe_style(self._doc, style_name, base)
        try:
            para = self._doc.add_paragraph(style=style) if style else self._doc.add_paragraph()
        except Exception:
            para = self._doc.add_paragraph()
        self._emit_runs(para, inlines)

    def _flush_code_block(self) -> None:
        """누적된 <pre> 원시 텍스트를 monospace 단락(들)로 내보낸다.

        renderer 의 코드블록은 ``<pre><code class="language-X">
        <div class="codehilite"><pre>…spans…</pre></div></code></pre>`` 이중 pre 구조다.
        ``convert_charrefs=True`` + span 무시(태그는 데이터가 아님)로 내부 텍스트만
        순서대로 이어붙이면 개행 포함 원본 코드가 복원된다.
        """
        code_text = "".join(self._pre_text)
        self._pre_text = []
        # 코드블록 끝의 잉여 개행 1개만 트림(원본 라인 보존).
        if code_text.endswith("\n"):
            code_text = code_text[:-1]
        if not code_text:
            return
        for line in code_text.split("\n"):
            try:
                para = self._doc.add_paragraph()
                run = para.add_run(line if line else "")
                run.font.name = _MONO_FONT
                run.font.size = _pt(_CODE_FONT_SIZE_PT)
                _set_run_shading(run, "F2F2F2")
            except Exception:
                try:
                    self._doc.add_paragraph(line)
                except Exception:
                    pass

    def _flush_table(self) -> None:
        info = self._table_stack.pop() if self._table_stack else None
        if not info:
            return
        rows = info.get("rows") or []
        cols = info.get("cols") or 0
        if not rows or cols == 0:
            return
        try:
            table = self._doc.add_table(rows=len(rows), cols=cols)
            style = _safe_style(self._doc, "Light Grid Accent 1", "Table Grid")
            if style:
                try:
                    table.style = style
                except Exception:
                    pass
            for r, row in enumerate(rows):
                for c in range(cols):
                    text = row[c] if c < len(row) else ""
                    try:
                        table.rows[r].cells[c].text = text
                    except Exception:
                        pass
        except Exception:
            # 표 생성 실패 시 텍스트라도 보존.
            for row in rows:
                try:
                    self._doc.add_paragraph("\t".join(row))
                except Exception:
                    pass

    # -- 자기완결 요소 -----------------------------------------------------
    def _handle_br(self) -> None:
        target = self._cell_inlines if self._cell_inlines is not None else self._inlines
        target.append(_Inline(is_break=True))

    def _handle_hr(self) -> None:
        # 구분선: "─" 반복 단락으로 간략 표현(설계 §2.3).
        try:
            self._doc.add_paragraph("─" * 30)  # ─
        except Exception:
            pass

    def _handle_input(self, attrs: dict) -> None:
        # task-list checkbox: 상위 li 의 체크 상태를 갱신.
        if self._li_task and self._li_task[-1] is not None:
            self._li_task[-1] = "checked" in attrs

    def _handle_img(self, attrs: dict) -> None:
        src = attrs.get("src") or ""
        alt = attrs.get("alt") or ""
        local_path = _file_uri_to_path(src) if src else None
        if local_path is not None and local_path.exists() and local_path.is_file():
            try:
                from docx.shared import Inches

                # 본문 폭 초과 방지: 임베드 후 폭 클램프는 add_picture width 로.
                # 우선 원본 폭으로 시도하되, 과대하면 본문 폭으로 클램프.
                pic = self._doc.add_picture(str(local_path))
                try:
                    if pic.width is not None and pic.width > Inches(_BODY_WIDTH_INCHES):
                        ratio = Inches(_BODY_WIDTH_INCHES) / pic.width
                        pic.width = Inches(_BODY_WIDTH_INCHES)
                        pic.height = int(pic.height * ratio)
                except Exception:
                    pass
                return
            except Exception:
                # 임베드 실패(포맷 미지원 등) → alt 폴백.
                pass
        # 원격/복원 실패/없음/임베드 실패 → alt(없으면 [이미지]) 이탤릭 run 폴백.
        fallback = alt if alt else "[이미지]"
        frag = _Inline(fallback, italic=True)
        if self._cell_inlines is not None:
            self._cell_inlines.append(frag)
        elif self._ctx:
            self._inlines.append(frag)
        else:
            # 블록 컨텍스트 밖(예: <p><img></p> 의 p 가 이미 닫혔거나) → 단독 단락.
            try:
                para = self._doc.add_paragraph()
                self._emit_runs(para, [frag])
            except Exception:
                pass


# ---------------------------------------------------------------------------
# 작은 유틸
# ---------------------------------------------------------------------------
def _pt(size: int):
    """Pt 객체 생성(지연 import)."""
    from docx.shared import Pt

    return Pt(size)


def _has_visible(inlines: list[_Inline]) -> bool:
    """인라인 조각 중 보이는 내용(공백 아닌 텍스트 또는 break)이 있는지."""
    for frag in inlines:
        if frag.is_break:
            return True
        if frag.text and frag.text.strip():
            return True
    return False


def _trim_boundary_ws(inlines: list[_Inline]) -> None:
    """리스트 항목 인라인 시퀀스의 **바깥쪽 공백**만 제자리(in place)에서 다듬는다.

    HTML 직렬화가 항목/자식목록 경계에 끼워 넣는 구조적 개행·들여쓰기(예:
    ``"Item B\\n"`` / ``"\\nNested"``)가 run 텍스트에 잔존하지 않도록, 첫 텍스트
    조각은 ``lstrip``, 마지막 텍스트 조각은 ``rstrip`` 한다. 내부 간격과 인라인
    서식(bold/italic/code/link/break)은 보존한다. break 조각은 트림 대상 아님.
    """
    if not inlines:
        return
    # 첫 텍스트 조각의 선행 공백 제거.
    for frag in inlines:
        if frag.is_break:
            break
        if frag.text:
            frag.text = frag.text.lstrip()
            break
    # 마지막 텍스트 조각의 후행 공백 제거.
    for frag in reversed(inlines):
        if frag.is_break:
            break
        if frag.text:
            frag.text = frag.text.rstrip()
            break


def _inline_plain(inlines: list[_Inline]) -> str:
    """인라인 조각들을 평문 문자열로 평탄화(헤딩/셀용). break 는 공백으로."""
    parts: list[str] = []
    for frag in inlines:
        if frag.is_break:
            parts.append(" ")
        elif frag.text:
            parts.append(frag.text)
    return "".join(parts).strip()


def _safe_style(doc: "_DocxDocument", name: str, fallback: str | None):
    """스타일이 템플릿에 있으면 그 이름을, 없으면 fallback(또는 None)을 반환한다."""
    try:
        styles = doc.styles
        if name in styles:
            return name
    except Exception:
        pass
    return fallback


def _file_uri_to_path(src: str) -> Path | None:
    """``file:///…`` URI 를 로컬 경로로 복원한다. 원격/비파일이면 None.

    Windows 드라이브 문자(``file:///D:/...``)를 올바르게 복원한다
    (``url2pathname(urlparse(src).path)``). 실패 시 None.
    """
    try:
        parts = urllib.parse.urlparse(src)
        if parts.scheme != "file":
            return None  # http(s)/data 등 원격 — 다운로드하지 않음.
        local = urllib.request.url2pathname(parts.path)
        if not local:
            return None
        return Path(local)
    except Exception:
        return None


# ---------------------------------------------------------------------------
# 공개 API
# ---------------------------------------------------------------------------
def markdown_to_docx(
    markdown_text: str,
    out_path: Path,
    base_dir: Path,
    *,
    title: str | None = None,
) -> None:
    """마크다운 텍스트를 Word(.docx) 문서로 내보낸다.

    내부적으로 ``renderer.render(markdown_text, base_dir)`` 로 본문 HTML 을 만든 뒤,
    표준 라이브러리 ``html.parser`` 로 walk 하며 python-docx 요소로 매핑한다
    (bs4/lxml 직접 파싱 금지 — 추가 의존성·번들 표면 회피).

    Args:
        markdown_text: 원본 마크다운 소스(``_doc_text``). None 이면 빈 문서로 취급.
        out_path: 저장 대상 ``.docx`` 경로. 부모 디렉터리가 없으면 생성한다.
        base_dir: 상대 이미지/링크 해석 기준 디렉터리(보통 열린 파일의 부모).
            renderer 가 상대경로를 ``file:///`` URI 로 치환할 때 사용하는 것과 동일.
        title: (선택) 문서 제목. 지정 시 docx 코어 속성(``core_properties.title``)에
            기록한다. 본문에 별도 제목 단락을 추가하지는 않는다(본문 h1 이 이미 제목).

    Returns:
        None.

    Raises:
        OSError: ``out_path`` 디렉터리 생성/파일 저장 실패(권한·디스크·잠금 등) 시
            **그대로 전파**한다(``write_markdown`` 과 대칭 — "저장 실패"는 사용자가
            반드시 알아야 하는 사건. UI 가 try/except 로 잡아 QMessageBox).
        그 외 예외는 던지지 않는다 — 아래 "예외 정책" 참조.

    예외 정책 (★ 두 철학의 결합 — 경계면 버그 단골):
        - **변환 자체는 절대 크래시하지 않는다**(``render`` 의 "어떤 입력에도 안전"
          철학). 깨진/빈/비정상 HTML, 알 수 없는 태그, 깨진 이미지에도 예외를 던지지
          않고 **최선의 docx 를 만들어 저장**한다. 매핑 불가 태그는 텍스트만 보존하거나
          건너뛴다. 빈 입력이면 빈(또는 사실상 빈) docx 를 저장한다.
        - **단, 최종 ``Document.save(out_path)`` 의 I/O 실패(OSError)만 전파**한다.
          변환 단계의 내부 오류는 삼키되, "파일을 못 썼다"는 전파한다.

    Thread-safety:
        순수 함수(전역 가변 상태 없음). 매 호출 새 ``docx.Document()`` 인스턴스 생성.
        ui-dev 는 메인 스레드에서 동기 호출한다(전형적 문서 크기에서 충분히 빠름).
        대용량 대비 워커 스레드 호출도 안전(인스턴스 비공유). v1 은 동기로 충분.
    """
    # python-docx(및 lxml)는 지연 import — 렌더 핫패스에서 lxml 로딩 회피(설계 §6.1).
    from docx import Document

    out_path = Path(out_path)

    # base_dir 정규화(renderer.render 와 동일한 방어).
    if not isinstance(base_dir, Path):
        try:
            base_dir = Path(base_dir)
        except (TypeError, ValueError):
            base_dir = Path.cwd()

    doc = Document()

    # 코어 속성(제목)만 기록 — 본문에는 별도 제목 단락을 추가하지 않는다.
    if title:
        try:
            doc.core_properties.title = title
        except Exception:
            pass

    # 본문 HTML 획득(render 는 예외를 던지지 않음). None 입력은 render 가 ""로 처리.
    try:
        result = render(markdown_text if markdown_text is not None else "", base_dir)
        body_html = result.html or ""
    except Exception:
        # render 가 절대 던지지 않지만, 만일을 대비해 변환 단계는 크래시 금지.
        body_html = ""

    # HTML walk → docx 매핑. 모든 변환 오류는 삼켜서 부분 결과라도 저장.
    if body_html.strip():
        builder = _DocxBuilder(doc, base_dir)
        try:
            builder.feed(body_html)
            builder.close()
        except Exception:
            # 파서 자체가 깨져도 변환은 크래시하지 않는다(부분 결과 보존).
            pass

    # 부모 디렉터리 보장(write_markdown 대칭). 실패 시 OSError 전파.
    out_path.parent.mkdir(parents=True, exist_ok=True)

    # 최종 저장 — OSError 만 전파(설계 §2.1).
    doc.save(str(out_path))
